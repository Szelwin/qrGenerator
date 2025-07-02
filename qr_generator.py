"""
qr_generator.py
Core utilities for producing QR codes and dropping them into a
Microsoft Word document (.docx).

You can import `create_qr_doc()` from another script (GUI, tests, etc.)
or treat this file as a CLI:

    python qr_generator.py 1000 1200

That would create "QR_1000_1200.docx".
"""

from __future__ import annotations

import argparse
from dataclasses import dataclass
from io import BytesIO
from typing import Union, Iterable, Optional
from pathlib import Path

import qrcode
from qrcode.constants import ERROR_CORRECT_L, ERROR_CORRECT_M, ERROR_CORRECT_Q, ERROR_CORRECT_H
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, Inches


@dataclass
class QRConfig:
    """Configuration for QR code generation."""
    box_size: int = 5
    border: int = 2
    error_correction: str = 'L'
    fill_color: str = 'black'
    back_color: str = 'white'
    version: Optional[int] = 1
    fit: bool = False
    
    def get_error_correction_constant(self):
        """Convert string error correction to qrcode constant."""
        mapping = {
            'L': ERROR_CORRECT_L,
            'M': ERROR_CORRECT_M, 
            'Q': ERROR_CORRECT_Q,
            'H': ERROR_CORRECT_H
        }
        return mapping.get(self.error_correction, ERROR_CORRECT_L)


@dataclass
class DocumentConfig:
    """Configuration for document formatting."""
    page_width_mm: float = 210.0  # A4 width
    page_height_mm: float = 297.0  # A4 height
    margin_inches: float = 0.5
    qr_width_mm: float = 9.0
    label_font_size_pt: float = 8.0
    columns: int = 17
    chunk_size: int = 100


# ──────────────────────────────────────────────────────────────────────────────
# Core QR generation functions
# ──────────────────────────────────────────────────────────────────────────────
def create_qr_png_stream(data: Union[int, str], config: QRConfig = None) -> BytesIO:
    """Return an in-memory PNG stream for a single QR code.
    
    Args:
        data: The data to encode in the QR code
        config: QR configuration options
        
    Returns:
        BytesIO: In-memory PNG stream of the QR code
    """
    if config is None:
        config = QRConfig()
    
    qr = qrcode.QRCode(
        version=config.version,
        error_correction=config.get_error_correction_constant(),
        box_size=config.box_size,
        border=config.border,
    )

    qr.add_data(data)
    qr.make(fit=config.fit)

    img = qr.make_image(fill_color=config.fill_color, back_color=config.back_color)
    buffer = BytesIO()

    img.save(buffer, format="PNG")
    buffer.seek(0)

    return buffer


def chunk_range(
    start: int, end_exclusive: int, size: int = 100
) -> Iterable[tuple[int, int]]:
    """
    Yield (chunk_start, chunk_end_inclusive) for `range(start, end_exclusive)`
    in blocks of `size`. Works even if the last chunk is < size.
    
    Args:
        start: Starting number (inclusive)
        end_exclusive: Ending number (exclusive)
        size: Chunk size
        
    Yields:
        tuple[int, int]: (chunk_start, chunk_end_inclusive) pairs
    """
    if start >= end_exclusive:
        return
        
    cur = start
    while cur < end_exclusive:
        yield cur, min(cur + size - 1, end_exclusive - 1)
        cur += size


def create_document(config: DocumentConfig = None) -> Document:
    """Return a fresh document with specified formatting.
    
    Args:
        config: Document configuration options
        
    Returns:
        Document: Configured Word document
    """
    if config is None:
        config = DocumentConfig()
        
    doc = Document()
    for section in doc.sections:
        section.page_width = Mm(config.page_width_mm)
        section.page_height = Mm(config.page_height_mm)
        margin = Inches(config.margin_inches)
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = margin
    return doc


# ──────────────────────────────────────────────────────────────────────────────
# Public API
# ──────────────────────────────────────────────────────────────────────────────
def add_qr_block(
    doc: Document, 
    start_num: int, 
    end_num: int, 
    qr_config: QRConfig = None,
    doc_config: DocumentConfig = None
) -> None:
    """Insert QR codes for start_num … end_num (inclusive) into doc.

    The codes are laid out in a single table with configurable columns. A range
    label (e.g. "101-200") is placed to the immediate right of the last code
    when room permits; otherwise it shares that cell.
    
    Args:
        doc: Word document to add QR codes to
        start_num: First number to encode (inclusive)
        end_num: Last number to encode (inclusive)
        qr_config: QR code generation configuration
        doc_config: Document formatting configuration
        
    Raises:
        ValueError: If end_num < start_num
    """
    if end_num < start_num:
        raise ValueError("end_num must be ≥ start_num")
    
    if qr_config is None:
        qr_config = QRConfig()
    if doc_config is None:
        doc_config = DocumentConfig()

    total = end_num - start_num + 1
    cols = doc_config.columns
    rows = -(-total // cols)  # ceiling division without math.ceil()

    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = False

    for idx, number in enumerate(range(start_num, end_num + 1)):
        r, c = divmod(idx, cols)
        cell = table.cell(r, c)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        qr_stream = create_qr_png_stream(number, qr_config)
        p.add_run().add_picture(qr_stream, width=Mm(doc_config.qr_width_mm))

    # Add trailing range label
    r, c = divmod(total - 1, cols)
    label_cell = table.cell(r, c + 1 if c < cols - 1 else c)
    label_para = label_cell.paragraphs[0]
    label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = label_para.add_run(f"{start_num}-{end_num}")
    run.font.size = Pt(doc_config.label_font_size_pt)

    # Add spacer lines after each block
    doc.add_paragraph()
    doc.add_paragraph()


def create_qr_doc(
    start: int, 
    end_exclusive: int, 
    out_path: str | Path,
    qr_config: QRConfig = None,
    doc_config: DocumentConfig = None
) -> Path:
    """
    Build a .docx file with QR codes for all numbers in range(start, end_exclusive).
    
    Args:
        start: First number to encode (inclusive)
        end_exclusive: Last number (exclusive)
        out_path: Output file path
        qr_config: QR code generation configuration
        doc_config: Document formatting configuration
        
    Returns:
        Path: Path of the created file
        
    Raises:
        ValueError: If start >= end_exclusive
    """
    if start >= end_exclusive:
        raise ValueError("start must be < end_exclusive")
        
    if qr_config is None:
        qr_config = QRConfig()
    if doc_config is None:
        doc_config = DocumentConfig()
    
    doc = create_document(doc_config)
    for chunk_start, chunk_end in chunk_range(start, end_exclusive, doc_config.chunk_size):
        add_qr_block(doc, chunk_start, chunk_end, qr_config, doc_config)

    out_path = Path(out_path)
    doc.save(out_path)
    return out_path


# ──────────────────────────────────────────────────────────────────────────────
# CLI entry-point – optional, keeps parity with your original script
# ──────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate QR-code .docx files.")
    parser.add_argument("start", type=int, help="First integer (inclusive).")
    parser.add_argument("end", type=int, help="Last integer (exclusive).")

    args = parser.parse_args()
    fname = f"QR_{args.start}_{args.end}.docx"

    create_qr_doc(args.start, args.end, fname)
    print(f"Saved → {fname}")
import argparse
import qrcode
from qrcode.constants import (  # noqa
    ERROR_CORRECT_L,
    ERROR_CORRECT_M,
    ERROR_CORRECT_Q,
    ERROR_CORRECT_H,
)

from docx import Document
from docx.shared import Inches, Mm, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH

import os

COLUMNS = 17

# Creates a QR code with the filename
def generate_qr_code(data: int, file_name):
    qr = qrcode.QRCode(
        version=1, error_correction=ERROR_CORRECT_L, box_size=5, border=2
    )

    qr.add_data(data)
    qr.make(fit=False)

    img = qr.make_image(fill="black", back_color="white")
    img.save(file_name)


def create_document():
    # Create a new Word document
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)  # Set page width to A4 size (210 mm)
    section.page_height = Mm(297)  # Set page height to A4 size (297 mm)

    # Adjust the page margins
    sections = doc.sections
    margin = Inches(0.5)
    for section in sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

    return doc

def write_hundred(start: int, doc: Document):
    end = start + 99
    total_qr_codes = end - start + 1
    # Calculate the number of rows needed
    num_rows = (total_qr_codes + COLUMNS - 1) // COLUMNS  # Ceiling division

    # Create a table with the calculated number of rows and columns
    table = doc.add_table(rows=num_rows, cols=COLUMNS)
    table.autofit = False  # Prevent automatic resizing

    qr_index = 0
    for i in range(start, end + 1):
        qr_file_name = f"qr_code_{i}.png"
        qr_data = i
        generate_qr_code(qr_data, qr_file_name)

        # Compute row and column indices
        row_idx = qr_index // COLUMNS
        col_idx = qr_index % COLUMNS

        # Insert the QR code into the table cell
        cell = table.cell(row_idx, col_idx)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(qr_file_name, width=Mm(9))

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Clean up the image file if needed
        os.remove(qr_file_name)

        qr_index += 1

    # At the end of the last row, next to the last QR code, not below it,
    # write the range of the QR codes in small font
    # After adding all QR codes, write the range of QR codes in small font
    # Determine the position to insert the text
    last_qr_index = qr_index - 1
    last_row_idx = last_qr_index // COLUMNS
    last_col_idx = last_qr_index % COLUMNS

    if last_col_idx < COLUMNS - 1:
        # There is space to the right of the last QR code
        text_cell = table.cell(last_row_idx, last_col_idx + 1)
        text_paragraph = text_cell.paragraphs[0]
        text_run = text_paragraph.add_run(f"{start}-{end}")
        text_run.font.size = Pt(8)  # Set font size to 8 pt
        text_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # No space to the right; add text next to the last QR code in the same cell
        cell = table.cell(last_row_idx, last_col_idx)
        paragraph = cell.paragraphs[0]
        # Add a new run for the text after the image
        text_run = paragraph.add_run(f" {start}-{end}")
        text_run.font.size = Pt(8)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # add 2 new lines
    doc.add_paragraph()
    doc.add_paragraph()
    

def create_qr_codes_in_word(start_range, end_range, doc_name):
    doc = create_document()

    number_of_hundreds = (end_range - start_range) // 100
    for i in range(number_of_hundreds):
        start = start_range + i * 100
        write_hundred(start, doc)

    # Save the Word document
    doc.save(doc_name)

# Parse command-line arguments
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate QR codes and insert them into a Word document.")
    parser.add_argument("start", type=int, help="Starting number for QR code generation (inclusive)")
    parser.add_argument("end", type=int, help="Ending number for QR code generation (exclusive)")

    args = parser.parse_args()

    # Call the function with the command-line arguments
    create_qr_codes_in_word(args.start, args.end, f"QR_{args.start}_{args.end}.docx")
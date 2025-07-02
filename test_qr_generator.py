"""
Unit tests for qr_generator module with comprehensive coverage.
"""

import unittest
import tempfile
import os
from io import BytesIO
from pathlib import Path
from unittest.mock import patch, MagicMock

import qrcode
from qrcode.constants import ERROR_CORRECT_L, ERROR_CORRECT_M, ERROR_CORRECT_Q, ERROR_CORRECT_H
import docx
from docx import Document
from PIL import Image

from qr_generator import (
    QRConfig, 
    DocumentConfig, 
    create_qr_png_stream, 
    chunk_range, 
    create_document, 
    add_qr_block, 
    create_qr_doc
)


class TestQRConfig(unittest.TestCase):
    """Test QRConfig dataclass and its methods."""
    
    def test_default_values(self):
        """Test default configuration values."""
        config = QRConfig()
        self.assertEqual(config.box_size, 5)
        self.assertEqual(config.border, 2)
        self.assertEqual(config.error_correction, 'L')
        self.assertEqual(config.fill_color, 'black')
        self.assertEqual(config.back_color, 'white')
        self.assertEqual(config.version, 1)
        self.assertFalse(config.fit)
    
    def test_custom_values(self):
        """Test custom configuration values."""
        config = QRConfig(
            box_size=10,
            border=4,
            error_correction='H',
            fill_color='blue',
            back_color='yellow',
            version=2,
            fit=True
        )
        self.assertEqual(config.box_size, 10)
        self.assertEqual(config.border, 4)
        self.assertEqual(config.error_correction, 'H')
        self.assertEqual(config.fill_color, 'blue')
        self.assertEqual(config.back_color, 'yellow')
        self.assertEqual(config.version, 2)
        self.assertTrue(config.fit)
    
    def test_error_correction_constants(self):
        """Test error correction constant mapping."""
        config = QRConfig()
        
        config.error_correction = 'L'
        self.assertEqual(config.get_error_correction_constant(), ERROR_CORRECT_L)
        
        config.error_correction = 'M'
        self.assertEqual(config.get_error_correction_constant(), ERROR_CORRECT_M)
        
        config.error_correction = 'Q'
        self.assertEqual(config.get_error_correction_constant(), ERROR_CORRECT_Q)
        
        config.error_correction = 'H'
        self.assertEqual(config.get_error_correction_constant(), ERROR_CORRECT_H)
    
    def test_invalid_error_correction(self):
        """Test invalid error correction defaults to L."""
        config = QRConfig(error_correction='INVALID')
        self.assertEqual(config.get_error_correction_constant(), ERROR_CORRECT_L)


class TestDocumentConfig(unittest.TestCase):
    """Test DocumentConfig dataclass."""
    
    def test_default_values(self):
        """Test default document configuration values."""
        config = DocumentConfig()
        self.assertEqual(config.page_width_mm, 210.0)
        self.assertEqual(config.page_height_mm, 297.0)
        self.assertEqual(config.margin_inches, 0.5)
        self.assertEqual(config.qr_width_mm, 9.0)
        self.assertEqual(config.label_font_size_pt, 8.0)
        self.assertEqual(config.columns, 17)
        self.assertEqual(config.chunk_size, 100)
    
    def test_custom_values(self):
        """Test custom document configuration values."""
        config = DocumentConfig(
            page_width_mm=216.0,
            page_height_mm=279.0,
            margin_inches=1.0,
            qr_width_mm=12.0,
            label_font_size_pt=10.0,
            columns=20,
            chunk_size=50
        )
        self.assertEqual(config.page_width_mm, 216.0)
        self.assertEqual(config.page_height_mm, 279.0)
        self.assertEqual(config.margin_inches, 1.0)
        self.assertEqual(config.qr_width_mm, 12.0)
        self.assertEqual(config.label_font_size_pt, 10.0)
        self.assertEqual(config.columns, 20)
        self.assertEqual(config.chunk_size, 50)


class TestCreateQRPngStream(unittest.TestCase):
    """Test create_qr_png_stream function."""
    
    def test_with_default_config(self):
        """Test QR generation with default configuration."""
        stream = create_qr_png_stream("test_data")
        self.assertIsInstance(stream, BytesIO)
        self.assertGreater(stream.getvalue().__len__(), 0)
        
        # Verify it's a valid PNG
        stream.seek(0)
        img = Image.open(stream)
        self.assertEqual(img.format, 'PNG')
    
    def test_with_custom_config(self):
        """Test QR generation with custom configuration."""
        config = QRConfig(box_size=10, border=1, error_correction='H')
        stream = create_qr_png_stream(123, config)
        self.assertIsInstance(stream, BytesIO)
        self.assertGreater(stream.getvalue().__len__(), 0)
    
    def test_with_string_data(self):
        """Test QR generation with string data."""
        stream = create_qr_png_stream("hello world")
        self.assertIsInstance(stream, BytesIO)
        self.assertGreater(stream.getvalue().__len__(), 0)
    
    def test_with_integer_data(self):
        """Test QR generation with integer data."""
        stream = create_qr_png_stream(42)
        self.assertIsInstance(stream, BytesIO)
        self.assertGreater(stream.getvalue().__len__(), 0)
    
    @patch('qr_generator.qrcode.QRCode')
    def test_qr_code_parameters(self, mock_qr_code):
        """Test that QRCode is initialized with correct parameters."""
        mock_qr = MagicMock()
        mock_qr_code.return_value = mock_qr
        mock_qr.make_image.return_value = MagicMock()
        
        config = QRConfig(box_size=8, border=3, error_correction='M', version=2, fit=True)
        create_qr_png_stream("test", config)
        
        mock_qr_code.assert_called_once_with(
            version=2,
            error_correction=ERROR_CORRECT_M,
            box_size=8,
            border=3
        )
        mock_qr.add_data.assert_called_once_with("test")
        mock_qr.make.assert_called_once_with(fit=True)


class TestChunkRange(unittest.TestCase):
    """Test chunk_range function."""
    
    def test_basic_chunking(self):
        """Test basic range chunking."""
        chunks = list(chunk_range(1, 11, 3))
        expected = [(1, 3), (4, 6), (7, 9), (10, 10)]
        self.assertEqual(chunks, expected)
    
    def test_exact_chunks(self):
        """Test range that divides evenly into chunks."""
        chunks = list(chunk_range(0, 10, 5))
        expected = [(0, 4), (5, 9)]
        self.assertEqual(chunks, expected)
    
    def test_single_chunk(self):
        """Test range smaller than chunk size."""
        chunks = list(chunk_range(5, 8, 10))
        expected = [(5, 7)]
        self.assertEqual(chunks, expected)
    
    def test_large_chunk_size(self):
        """Test chunk size larger than range."""
        chunks = list(chunk_range(1, 5, 100))
        expected = [(1, 4)]
        self.assertEqual(chunks, expected)
    
    def test_empty_range(self):
        """Test empty range."""
        chunks = list(chunk_range(5, 5, 10))
        self.assertEqual(chunks, [])
    
    def test_invalid_range(self):
        """Test invalid range (start >= end)."""
        chunks = list(chunk_range(10, 5, 10))
        self.assertEqual(chunks, [])
    
    def test_single_item_range(self):
        """Test range with single item."""
        chunks = list(chunk_range(1, 2, 5))
        expected = [(1, 1)]
        self.assertEqual(chunks, expected)
    
    def test_default_chunk_size(self):
        """Test default chunk size of 100."""
        chunks = list(chunk_range(1, 101))
        expected = [(1, 100)]
        self.assertEqual(chunks, expected)
        
        chunks = list(chunk_range(1, 201))
        expected = [(1, 100), (101, 200)]
        self.assertEqual(chunks, expected)


class TestCreateDocument(unittest.TestCase):
    """Test create_document function."""
    
    def test_default_document(self):
        """Test document creation with default configuration."""
        doc = create_document()
        self.assertTrue(hasattr(doc, 'sections'))
        
        section = doc.sections[0]
        # Check page dimensions (allow for small floating point differences)
        self.assertAlmostEqual(section.page_width.mm, 210.0, places=1)
        self.assertAlmostEqual(section.page_height.mm, 297.0, places=1)
        self.assertEqual(section.top_margin.inches, 0.5)
    
    def test_custom_document(self):
        """Test document creation with custom configuration."""
        config = DocumentConfig(
            page_width_mm=216.0,
            page_height_mm=279.0,
            margin_inches=1.0
        )
        doc = create_document(config)
        self.assertTrue(hasattr(doc, 'sections'))
        
        section = doc.sections[0]
        self.assertAlmostEqual(section.page_width.mm, 216.0, places=1)
        self.assertAlmostEqual(section.page_height.mm, 279.0, places=1)
        self.assertEqual(section.top_margin.inches, 1.0)


class TestAddQRBlock(unittest.TestCase):
    """Test add_qr_block function."""
    
    def setUp(self):
        """Set up test fixtures."""
        self.doc = create_document()
    
    def test_basic_qr_block(self):
        """Test adding a basic QR block."""
        initial_paragraphs = len(self.doc.paragraphs)
        add_qr_block(self.doc, 1, 5)
        
        # Check that tables and paragraphs were added
        self.assertEqual(len(self.doc.tables), 1)
        self.assertGreater(len(self.doc.paragraphs), initial_paragraphs)
        
        table = self.doc.tables[0]
        self.assertEqual(table.rows.__len__(), 1)  # 5 items in 17 columns = 1 row
        self.assertEqual(table.columns.__len__(), 17)
    
    def test_multiple_rows(self):
        """Test QR block that spans multiple rows."""
        add_qr_block(self.doc, 1, 20)  # 20 items > 17 columns
        
        table = self.doc.tables[0]
        expected_rows = 2  # ceil(20/17) = 2
        self.assertEqual(table.rows.__len__(), expected_rows)
    
    def test_custom_configurations(self):
        """Test QR block with custom configurations."""
        qr_config = QRConfig(box_size=8, error_correction='H')
        doc_config = DocumentConfig(columns=5, qr_width_mm=15.0)
        
        add_qr_block(self.doc, 1, 10, qr_config, doc_config)
        
        table = self.doc.tables[0]
        expected_rows = 2  # ceil(10/5) = 2
        self.assertEqual(table.rows.__len__(), expected_rows)
        self.assertEqual(table.columns.__len__(), 5)
    
    def test_invalid_range(self):
        """Test error handling for invalid range."""
        with self.assertRaises(ValueError) as context:
            add_qr_block(self.doc, 10, 5)
        self.assertIn("end_num must be ≥ start_num", str(context.exception))
    
    def test_single_qr_code(self):
        """Test adding a single QR code."""
        add_qr_block(self.doc, 42, 42)
        
        table = self.doc.tables[0]
        self.assertEqual(table.rows.__len__(), 1)
        # Should have content in first cell
        first_cell = table.cell(0, 0)
        self.assertTrue(len(first_cell.paragraphs[0].runs) > 0)
    
    @patch('qr_generator.create_qr_png_stream')
    def test_qr_generation_called(self, mock_create_qr):
        """Test that QR generation is called for each number."""
        # Create a proper PNG stream
        from PIL import Image
        import io
        
        # Create a 1x1 pixel PNG
        img = Image.new('RGB', (1, 1), color='white')
        png_buffer = io.BytesIO()
        img.save(png_buffer, format='PNG')
        png_buffer.seek(0)
        
        mock_create_qr.return_value = png_buffer
        
        add_qr_block(self.doc, 5, 7)  # 3 numbers: 5, 6, 7
        
        self.assertEqual(mock_create_qr.call_count, 3)
        # Check that it was called with the right numbers
        call_args = [call[0][0] for call in mock_create_qr.call_args_list]
        self.assertEqual(call_args, [5, 6, 7])


class TestCreateQRDoc(unittest.TestCase):
    """Test create_qr_doc function."""
    
    def setUp(self):
        """Set up test fixtures."""
        self.temp_dir = tempfile.mkdtemp()
        self.test_file = Path(self.temp_dir) / "test_qr.docx"
    
    def tearDown(self):
        """Clean up test fixtures."""
        if self.test_file.exists():
            self.test_file.unlink()
        os.rmdir(self.temp_dir)
    
    def test_basic_document_creation(self):
        """Test basic document creation and saving."""
        result_path = create_qr_doc(1, 10, self.test_file)
        
        self.assertEqual(result_path, self.test_file)
        self.assertTrue(self.test_file.exists())
        self.assertGreater(self.test_file.stat().st_size, 0)
    
    def test_with_custom_configs(self):
        """Test document creation with custom configurations."""
        qr_config = QRConfig(box_size=3, border=1)
        doc_config = DocumentConfig(chunk_size=5, columns=10)
        
        result_path = create_qr_doc(1, 15, self.test_file, qr_config, doc_config)
        
        self.assertEqual(result_path, self.test_file)
        self.assertTrue(self.test_file.exists())
    
    def test_string_path(self):
        """Test with string path instead of Path object."""
        result_path = create_qr_doc(1, 5, str(self.test_file))
        
        self.assertEqual(result_path, self.test_file)
        self.assertTrue(self.test_file.exists())
    
    def test_invalid_range(self):
        """Test error handling for invalid range."""
        with self.assertRaises(ValueError) as context:
            create_qr_doc(10, 5, self.test_file)
        self.assertIn("start must be < end_exclusive", str(context.exception))
    
    def test_equal_start_end(self):
        """Test error handling when start equals end."""
        with self.assertRaises(ValueError):
            create_qr_doc(5, 5, self.test_file)
    
    @patch('qr_generator.add_qr_block')
    @patch('qr_generator.chunk_range')
    def test_chunking_behavior(self, mock_chunk_range, mock_add_qr_block):
        """Test that chunking works correctly."""
        mock_chunk_range.return_value = [(1, 50), (51, 100), (101, 120)]
        
        create_qr_doc(1, 121, self.test_file)
        
        # Verify chunking was called correctly
        mock_chunk_range.assert_called_once_with(1, 121, 100)  # default chunk size
        
        # Verify add_qr_block was called for each chunk
        self.assertEqual(mock_add_qr_block.call_count, 3)
        call_args = [call[0][1:3] for call in mock_add_qr_block.call_args_list]
        expected_args = [(1, 50), (51, 100), (101, 120)]
        self.assertEqual(call_args, expected_args)
    
    def test_large_range(self):
        """Test handling of large ranges that span multiple chunks."""
        # This tests the actual chunking behavior
        result_path = create_qr_doc(1, 250, self.test_file)  # 249 QR codes, 3 chunks
        
        self.assertTrue(self.test_file.exists())
        
        # Load the document to verify structure
        doc = Document(str(self.test_file))
        # Should have 3 tables (3 chunks) plus spacing paragraphs
        self.assertEqual(len(doc.tables), 3)


class TestIntegration(unittest.TestCase):
    """Integration tests for the complete workflow."""
    
    def setUp(self):
        """Set up test fixtures."""
        self.temp_dir = tempfile.mkdtemp()
        self.test_file = Path(self.temp_dir) / "integration_test.docx"
    
    def tearDown(self):
        """Clean up test fixtures."""
        if self.test_file.exists():
            self.test_file.unlink()
        os.rmdir(self.temp_dir)
    
    def test_end_to_end_workflow(self):
        """Test complete end-to-end workflow."""
        # Custom configurations
        qr_config = QRConfig(
            box_size=4,
            border=1,
            error_correction='M',
            fill_color='darkblue',
            back_color='lightgray'
        )
        
        doc_config = DocumentConfig(
            page_width_mm=216.0,  # US Letter width
            page_height_mm=279.0,  # US Letter height
            margin_inches=0.75,
            qr_width_mm=8.0,
            label_font_size_pt=9.0,
            columns=15,
            chunk_size=75
        )
        
        # Create document
        result_path = create_qr_doc(100, 200, self.test_file, qr_config, doc_config)
        
        # Verify file creation
        self.assertTrue(self.test_file.exists())
        self.assertGreater(self.test_file.stat().st_size, 10000)  # Reasonable size
        
        # Load and verify document structure
        doc = Document(str(self.test_file))
        
        # Should have 2 chunks: 100-174 (75 items) and 175-199 (25 items)
        self.assertEqual(len(doc.tables), 2)
        
        # Verify first table dimensions
        first_table = doc.tables[0]
        expected_rows_first = 5  # ceil(75/15) = 5
        self.assertEqual(first_table.rows.__len__(), expected_rows_first)
        self.assertEqual(first_table.columns.__len__(), 15)
        
        # Verify second table dimensions
        second_table = doc.tables[1]
        expected_rows_second = 2  # ceil(25/15) = 2
        self.assertEqual(second_table.rows.__len__(), expected_rows_second)
    
    def test_minimal_configuration(self):
        """Test with minimal/default configuration."""
        result_path = create_qr_doc(1, 3, self.test_file)
        
        self.assertTrue(self.test_file.exists())
        
        # Load and verify
        doc = Document(str(self.test_file))
        self.assertEqual(len(doc.tables), 1)
        
        table = doc.tables[0]
        self.assertEqual(table.rows.__len__(), 1)  # 2 QR codes in 17 columns = 1 row
        self.assertEqual(table.columns.__len__(), 17)


if __name__ == '__main__':
    # Run with coverage if available
    try:
        import coverage
        cov = coverage.Coverage()
        cov.start()
        
        unittest.main(exit=False)
        
        cov.stop()
        cov.save()
        print("\nCoverage Report:")
        cov.report(show_missing=True)
    except ImportError:
        # Fall back to regular unittest
        unittest.main()